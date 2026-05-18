"""FastAPI server — HTTP interface for the Tauri frontend."""

from __future__ import annotations
import json
import os
import shutil
import time
import stat
from datetime import date, timedelta
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from ..agent import AgentSession
from ..llm.registry import PROVIDERS, PROVIDER_DEFAULTS, create_provider
from ..tools.registry import TOOLS
from ..sandbox.executor import execute_code
from ..data.cache import get_cached_preview, set_cached_preview
from ..data.index import get_index, invalidate_index


app = FastAPI(title="SheetAgent API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------------------------
# Config persistence (JSON file) - user directory
# ---------------------------------------------------------------------------
from datetime import date

_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent


def _get_user_config_dir() -> Path:
    """Get user config directory: ~/.sheet-agent/"""
    home = Path.home()
    config_dir = home / ".sheet-agent"
    config_dir.mkdir(exist_ok=True)
    return config_dir


def _get_default_workspace_root() -> Path:
    """Get default workspace root: ~/SheetAgent/workspace/"""
    home = Path.home()
    ws_root = home / "SheetAgent" / "workspace"
    ws_root.mkdir(parents=True, exist_ok=True)
    return ws_root


_CONFIG_FILE = _get_user_config_dir() / "settings.json"
_INSTALL_CONFIG_FILE = _PROJECT_ROOT / "config.yaml"
_WORKSPACE_ROOT = _get_default_workspace_root()


def _default_workspace() -> str:
    """Create workspace/YYYY-MM-DD if it doesn't exist, return path."""
    today = date.today().isoformat()
    ws = _WORKSPACE_ROOT / today
    ws.mkdir(parents=True, exist_ok=True)
    return str(ws)


_DEFAULT_CONFIG = {
    "provider": "claude",
    "model": "",
    "api_key": "",
    "base_url": "",
    "workspace": _default_workspace(),
}


def _load_yaml_config() -> dict:
    """Load config from install directory config.yaml if exists."""
    if not _INSTALL_CONFIG_FILE.exists():
        return {}

    try:
        import yaml
        with open(_INSTALL_CONFIG_FILE, encoding="utf-8") as f:
            data = yaml.safe_load(f) or {}

        # Flatten nested config structure
        result = {}
        if "llm" in data:
            result["provider"] = data["llm"].get("provider", "")
            result["model"] = data["llm"].get("model", "")
            result["api_key"] = data["llm"].get("api_key", "")
            result["base_url"] = data["llm"].get("base_url", "")
        if "workspace" in data:
            ws = data["workspace"]
            if isinstance(ws, dict):
                result["workspace"] = ws.get("default_path", "")
            else:
                result["workspace"] = ws

        # Filter empty values
        return {k: v for k, v in result.items() if v}
    except Exception as e:
        print(f"[Config] Failed to load install config: {e}")
        return {}


def _migrate_old_config():
    """Migrate config from old location (project_dir/config/) to new location (~/.sheet-agent/)."""
    old_config_file = _PROJECT_ROOT / "config" / "settings.json"
    if old_config_file.exists() and not _CONFIG_FILE.exists():
        try:
            old_cfg = json.loads(old_config_file.read_text(encoding="utf-8"))
            # Save to new location
            _CONFIG_FILE.parent.mkdir(parents=True, exist_ok=True)
            _CONFIG_FILE.write_text(json.dumps(old_cfg, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"[Config] Migrated config from {old_config_file} to {_CONFIG_FILE}")
        except (json.JSONDecodeError, OSError) as e:
            print(f"[Config] Failed to migrate old config: {e}")


def _resolve_workspace_path(workspace: str) -> str:
    """Resolve workspace path - support both relative and absolute paths.

    - Absolute path: use as-is
    - Relative path: resolve relative to user's default workspace root
    - Legacy path (under project root): migrate to new location
    """
    if not workspace:
        return _default_workspace()

    ws_path = Path(workspace)

    # If it's already an absolute path, check if it exists or create it
    if ws_path.is_absolute():
        try:
            ws_path.mkdir(parents=True, exist_ok=True)
        except OSError:
            pass
        return str(ws_path)

    # Check if it's a legacy path under project root
    old_style_path = _PROJECT_ROOT / "workspace" / workspace
    if old_style_path.exists():
        # Migrate to new location
        new_path = _WORKSPACE_ROOT / workspace
        try:
            new_path.parent.mkdir(parents=True, exist_ok=True)
            # Note: We don't move files automatically to avoid data loss
            # Just update the path to new location
            return str(new_path)
        except OSError:
            pass

    # Relative path - resolve under default workspace root
    resolved = _WORKSPACE_ROOT / workspace
    resolved.mkdir(parents=True, exist_ok=True)
    return str(resolved)


def _load_config() -> dict:
    """Load config with priority: user config > install dir config > defaults."""
    # Migrate old config if exists
    _migrate_old_config()

    # Start with defaults
    cfg = {**_DEFAULT_CONFIG}

    # Load from install directory (lower priority)
    install_cfg = _load_yaml_config()
    if install_cfg:
        cfg.update(install_cfg)

    # Load from user directory (highest priority)
    if _CONFIG_FILE.exists():
        try:
            saved = json.loads(_CONFIG_FILE.read_text(encoding="utf-8"))
            cfg.update(saved)
        except (json.JSONDecodeError, OSError):
            pass

    # Resolve workspace path
    cfg["workspace"] = _resolve_workspace_path(cfg.get("workspace"))
    return cfg


def _save_config(cfg: dict):
    _CONFIG_FILE.parent.mkdir(parents=True, exist_ok=True)
    # Write without api_key masking
    _CONFIG_FILE.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")


_sessions: dict[str, AgentSession] = {}
_config: dict = _load_config()

# Ensure workspace directory exists on startup
Path(_config["workspace"]).mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Session persistence
# ---------------------------------------------------------------------------

_SESSIONS_DIR = Path(_config["workspace"]) / ".sessions"
_SESSIONS_DIR.mkdir(parents=True, exist_ok=True)


def _session_file(session_id: str) -> Path:
    return _SESSIONS_DIR / f"{session_id}.json"


def _save_session(session: AgentSession):
    """Persist session state to disk."""
    try:
        data = session.to_dict()
        _session_file(session.session_id).write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass


def _load_session(session_id: str) -> AgentSession | None:
    """Restore session from disk if available."""
    sf = _session_file(session_id)
    if not sf.exists():
        return None
    try:
        data = json.loads(sf.read_text(encoding="utf-8"))
        return AgentSession.from_dict(data)
    except Exception:
        return None


def _list_saved_sessions() -> list[dict]:
    """List all persisted sessions with metadata."""
    sessions = []
    if not _SESSIONS_DIR.exists():
        return sessions
    for f in sorted(_SESSIONS_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            msgs = data.get("messages", [])
            # Find first user message for title
            title = "Untitled"
            for m in msgs:
                if m.get("role") == "user":
                    title = m.get("content", "Untitled")[:60]
                    break

            # Determine status
            session_id = data.get("session_id", f.stem)
            status = "success"
            if session_id in _sessions:
                status = "active"
            elif msgs:
                last_msg = msgs[-1]
                if last_msg.get("role") == "assistant":
                    text = last_msg.get("text", "")
                    if "error" in text.lower() or last_msg.get("error"):
                        status = "error"

            sessions.append({
                "session_id": session_id,
                "title": title,
                "turn_count": data.get("turn_count", 0),
                "updated_at": f.stat().st_mtime,
                "status": status,
            })
        except Exception:
            pass
    return sessions


def _delete_saved_session(session_id: str):
    sf = _session_file(session_id)
    if sf.exists():
        sf.unlink()


def _cleanup_old_tmp_dirs():
    """Remove .tmp subdirectories older than 7 days inside the workspace."""
    ws = Path(_config["workspace"])
    tmp_root = ws / ".tmp"
    if not tmp_root.exists():
        return
    cutoff = date.today() - timedelta(days=7)
    for d in tmp_root.iterdir():
        if d.is_dir():
            try:
                d_date = date.fromisoformat(d.name)
                if d_date < cutoff:
                    shutil.rmtree(d)
            except ValueError:
                pass


_cleanup_old_tmp_dirs()


# ---------------------------------------------------------------------------
# Request / Response models
# ---------------------------------------------------------------------------

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    stream: bool = True
    current_dir: Optional[str] = None
    lang: Optional[str] = None

class ConfigRequest(BaseModel):
    provider: Optional[str] = None
    model: Optional[str] = None
    api_key: Optional[str] = None
    base_url: Optional[str] = None
    workspace: Optional[str] = None

class ExecuteCodeRequest(BaseModel):
    code: str
    workspace: Optional[str] = None


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@app.get("/api/providers")
def list_providers():
    """List available LLM providers and their defaults."""
    return {
        "providers": [
            {
                "name": name,
                "default_model": defaults.get("model", ""),
                "default_base_url": defaults.get("base_url"),
            }
            for name, defaults in PROVIDER_DEFAULTS.items()
        ]
    }


@app.get("/api/config")
def get_config():
    """Get current configuration (api_key masked)."""
    cfg = {**_config}
    if cfg.get("api_key"):
        key = cfg["api_key"]
        cfg["api_key"] = key[:4] + "****" + key[-4:] if len(key) > 8 else "****"
    return cfg


@app.post("/api/config")
def update_config(req: ConfigRequest):
    """Update configuration."""
    if req.provider is not None:
        if req.provider not in PROVIDERS:
            raise HTTPException(400, f"Unknown provider: {req.provider}")
        _config["provider"] = req.provider
        # Reset model to default if provider changed
        if req.model is None:
            _config["model"] = PROVIDER_DEFAULTS[req.provider].get("model", "")

    if req.model is not None:
        _config["model"] = req.model
    if req.api_key is not None:
        _config["api_key"] = req.api_key
    if req.base_url is not None:
        _config["base_url"] = req.base_url
    if req.workspace is not None:
        # 使用 _resolve_workspace_path 统一处理路径，确保绝对路径正确保存
        resolved = _resolve_workspace_path(req.workspace)
        _config["workspace"] = resolved

    # Invalidate existing sessions so they pick up new config
    _sessions.clear()

    # Persist to disk
    _save_config(_config)

    return {"status": "ok"}


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), dir: str = ""):
    """Upload a file to the current workspace (optionally into a subdirectory)."""
    ws = Path(_config["workspace"])
    target_dir = ws / dir if dir else ws
    target_dir.mkdir(parents=True, exist_ok=True)

    if not str(target_dir.resolve()).startswith(str(ws.resolve())):
        raise HTTPException(403, "Access denied: path outside workspace")

    dest = target_dir / file.filename
    if dest.exists():
        # Don't overwrite — add suffix
        stem = dest.stem
        suffix = dest.suffix
        i = 1
        while dest.exists():
            dest = target_dir / f"{stem}_{i}{suffix}"
            i += 1

    content = await file.read()
    dest.write_bytes(content)
    invalidate_index(str(ws))

    return {
        "filename": dest.name,
        "path": str(dest.relative_to(ws)),
        "size_mb": round(len(content) / 1024 / 1024, 2),
    }


@app.get("/api/workspace/files")
def list_workspace_files(dir: str = "", pattern: str = "*"):
    """List files and folders in a workspace subdirectory."""
    ws = Path(_config["workspace"])
    target = ws / dir if dir else ws
    if not target.exists():
        return {"files": [], "dirs": [], "count": 0, "path": dir}

    # Security check
    if not str(target.resolve()).startswith(str(ws.resolve())):
        raise HTTPException(403, "Access denied: path outside workspace")

    items = sorted(target.iterdir())
    files = []
    dirs = []
    for f in items:
        rel = str(f.relative_to(ws))
        if f.is_dir():
            dirs.append({"name": f.name, "path": rel})
        elif f.is_file():
            stat = f.stat()
            files.append({
                "name": f.name,
                "path": rel,
                "size_mb": round(stat.st_size / 1024 / 1024, 2),
            })
    return {"files": files, "dirs": dirs, "count": len(files) + len(dirs), "path": dir}


@app.get("/api/workspace/stats")
def workspace_stats():
    """Return aggregate statistics for the workspace."""
    ws = _config["workspace"]
    idx = get_index(ws)
    snap = idx.get()
    ext_counts = {ext: len(entries) for ext, entries in snap.by_ext.items()}
    return {
        "workspace": ws,
        "total_files": snap.total_files,
        "total_size_mb": round(snap.total_bytes / 1024 / 1024, 2),
        "by_extension": ext_counts,
        "scanned_at": snap.scanned_at,
    }


@app.post("/api/workspace/mkdir")
def create_folder(req: dict):
    """Create a folder inside the workspace."""
    ws = Path(_config["workspace"])
    rel = req.get("dir", "")
    name = req.get("name", "").strip()
    if not name:
        raise HTTPException(400, "Folder name is required")
    target = (ws / rel / name) if rel else (ws / name)
    if not str(target.resolve()).startswith(str(ws.resolve())):
        raise HTTPException(403, "Access denied: path outside workspace")
    target.mkdir(parents=True, exist_ok=True)
    return {"status": "ok", "path": str(target.relative_to(ws))}


@app.post("/api/chat")
async def chat(req: ChatRequest):
    """Chat with the agent. Returns SSE stream if stream=True."""

    # Validate config
    if not _config["api_key"] and _config["provider"] != "ollama":
        raise HTTPException(400, "API key not configured. POST /api/config first.")

    workspace = _config["workspace"]
    if not Path(workspace).exists():
        raise HTTPException(400, f"Workspace does not exist: {workspace}")

    # Get or create session
    current_dir = req.current_dir or ""
    req_lang = req.lang or "en"
    if req.session_id and req.session_id in _sessions:
        session = _sessions[req.session_id]
        session.current_dir = current_dir
    elif req.session_id:
        # Try to restore from disk
        session = _load_session(req.session_id)
        if session:
            session.current_dir = current_dir
            _sessions[session.session_id] = session
        else:
            session = AgentSession.create(
                provider_name=_config["provider"],
                workspace=workspace,
                current_dir=current_dir,
                api_key=_config["api_key"] or None,
                model=_config["model"] or None,
                base_url=_config["base_url"] or None,
                lang=req_lang,
            )
            _sessions[session.session_id] = session
    else:
        session = AgentSession.create(
            provider_name=_config["provider"],
            workspace=workspace,
            current_dir=current_dir,
            api_key=_config["api_key"] or None,
            model=_config["model"] or None,
            base_url=_config["base_url"] or None,
            lang=req_lang,
        )
        _sessions[session.session_id] = session

    if req.stream:
        return StreamingResponse(
            _stream_response(session, req.message),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
            },
        )
    else:
        result = session.run(req.message)
        _save_session(session)
        return {
            "session_id": session.session_id,
            "turn_count": session.turn_count,
            "content": result,
        }


def _stream_response(session: AgentSession, message: str):
    """SSE generator for streaming chat responses."""
    try:
        yield f"data: {json.dumps({'type': 'session', 'session_id': session.session_id})}\n\n"

        for event in session.run_stream(message):
            yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

        yield "data: [DONE]\n\n"
    except Exception as e:
        # 发送错误事件给前端，而不是让流中断
        yield f"data: {json.dumps({'type': 'error', 'content': str(e)}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"
    finally:
        _save_session(session)


@app.post("/api/execute")
def execute_code_endpoint(req: ExecuteCodeRequest):
    """Directly execute Python code (for manual code editing in GUI)."""
    workspace = req.workspace or _config["workspace"]
    result = execute_code(req.code, workspace)
    return {
        "success": result.success,
        "stdout": result.stdout,
        "stderr": result.stderr,
        "exit_code": result.exit_code,
        "duration_ms": result.duration_ms,
    }


@app.get("/api/tools")
def list_tools():
    """List available tools."""
    return {
        "tools": [
            {"name": t.name, "description": t.description, "parameters": t.parameters}
            for t in TOOLS
        ]
    }


@app.get("/api/sessions")
def list_sessions():
    """List all persisted sessions."""
    return {"sessions": _list_saved_sessions()}


@app.get("/api/sessions/{session_id}")
def get_session(session_id: str):
    """Get a single session's conversation history."""
    # Check memory first
    if session_id in _sessions:
        session = _sessions[session_id]
        return {
            "session_id": session.session_id,
            "turn_count": session.turn_count,
            "messages": session.messages,
        }
    # Fallback to disk
    session = _load_session(session_id)
    if session:
        _sessions[session_id] = session
        return {
            "session_id": session.session_id,
            "turn_count": session.turn_count,
            "messages": session.messages,
        }
    raise HTTPException(404, "Session not found")


@app.delete("/api/sessions/{session_id}")
def delete_session(session_id: str):
    """Delete a session from memory and disk."""
    if session_id in _sessions:
        del _sessions[session_id]
    _delete_saved_session(session_id)
    return {"status": "ok"}


@app.get("/api/health")
def health():
    return {"status": "ok", "version": "0.1.0"}


# ---------------------------------------------------------------------------
# Data preview — read a file and return tabular data for the frontend
# ---------------------------------------------------------------------------

@app.get("/api/preview")
def preview_file(path: str, n_rows: int = 100):
    """Preview a file in the workspace as tabular data."""
    import polars as pl

    full_path = Path(_config["workspace"]) / path
    if not full_path.exists():
        raise HTTPException(404, f"File not found: {path}")
    if not str(full_path.resolve()).startswith(str(Path(_config["workspace"]).resolve())):
        raise HTTPException(403, "Access denied: path outside workspace")

    # Try cache first for tabular / document previews
    cached = get_cached_preview(_config["workspace"], str(full_path), n_rows)
    if cached is not None:
        return cached

    try:
        suffix = full_path.suffix.lower().lstrip(".")
        if suffix == "csv":
            df = pl.read_csv(str(full_path), infer_schema_length=200, ignore_errors=True)
            n = min(df.height, n_rows)
            preview = df.head(n)
            result = {
                "columns": preview.columns,
                "dtypes": [str(d) for d in preview.dtypes],
                "row_count": df.height,
                "preview_rows": n,
                "data": preview.to_dicts(),
            }
            set_cached_preview(_config["workspace"], str(full_path), n_rows, result)
            return result
        elif suffix == "parquet":
            df = pl.read_parquet(str(full_path))
            n = min(df.height, n_rows)
            preview = df.head(n)
            result = {
                "columns": preview.columns,
                "dtypes": [str(d) for d in preview.dtypes],
                "row_count": df.height,
                "preview_rows": n,
                "data": preview.to_dicts(),
            }
            set_cached_preview(_config["workspace"], str(full_path), n_rows, result)
            return result
        elif suffix in ("xlsx", "xls"):
            # 多 sheet 支持：先获取 sheet 名称列表，再逐个读取
            sheet_names = None
            try:
                import openpyxl
                wb = openpyxl.load_workbook(str(full_path), read_only=True, data_only=True)
                sheet_names = wb.sheetnames
            except Exception:
                pass

            if sheet_names:
                sheets = {}
                active_sheet = None
                for idx, name in enumerate(sheet_names, start=1):
                    try:
                        df = pl.read_excel(str(full_path), sheet_id=idx)
                    except Exception:
                        continue
                    if df.height == 0:
                        continue
                    if active_sheet is None:
                        active_sheet = name
                    n = min(df.height, n_rows)
                    preview = df.head(n)
                    sheets[name] = {
                        "columns": preview.columns,
                        "dtypes": [str(d) for d in preview.dtypes],
                        "row_count": df.height,
                        "preview_rows": n,
                        "data": preview.to_dicts(),
                    }
                if not sheets:
                    raise HTTPException(400, "All sheets are empty")
                if len(sheets) == 1:
                    # 只有一个有数据 sheet，按旧格式返回（兼容前端无 tab 展示）
                    result = list(sheets.values())[0]
                    set_cached_preview(_config["workspace"], str(full_path), n_rows, result)
                    return result
                result = {"sheets": sheets, "active_sheet": active_sheet}
                set_cached_preview(_config["workspace"], str(full_path), n_rows, result)
                return result
            else:
                # fallback：直接读取第一个 sheet
                try:
                    df = pl.read_excel(str(full_path), sheet_id=1)
                except Exception as e:
                    raise HTTPException(400, str(e)[:500])
                if df.height == 0:
                    raise HTTPException(400, "Empty Excel sheet")
                n = min(df.height, n_rows)
                preview = df.head(n)
                result = {
                    "columns": preview.columns,
                    "dtypes": [str(d) for d in preview.dtypes],
                    "row_count": df.height,
                    "preview_rows": n,
                    "data": preview.to_dicts(),
                }
                set_cached_preview(_config["workspace"], str(full_path), n_rows, result)
                return result
        elif suffix in ("docx", "doc"):
            # Word document preview
            text_preview = ""
            tables = []
            paragraph_count = 0
            if suffix == "docx":
                try:
                    from docx import Document
                    doc = Document(str(full_path))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    paragraph_count = len(paragraphs)
                    text_preview = "\n".join(paragraphs[:200])
                    for table in doc.tables[:5]:
                        tbl_data = []
                        for row in table.rows:
                            tbl_data.append([cell.text for cell in row.cells])
                        tables.append(tbl_data)
                except Exception as e:
                    text_preview = f"Error reading document: {str(e)}"
            else:
                text_preview = ".doc preview is not supported. Please convert to .docx for preview."
            result = {
                "type": "document",
                "format": suffix,
                "text_preview": text_preview[:3000],
                "paragraph_count": paragraph_count,
                "tables": tables,
            }
            set_cached_preview(_config["workspace"], str(full_path), n_rows, result)
            return result
        elif suffix in ("png", "jpg", "jpeg", "gif", "webp", "bmp"):
            # Image preview: return a data URL or relative path for frontend
            import base64
            data = full_path.read_bytes()
            mime = {
                "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "gif": "image/gif", "webp": "image/webp", "bmp": "image/bmp",
            }.get(suffix, "image/png")
            b64 = base64.b64encode(data).decode("utf-8")
            return {
                "type": "image",
                "format": suffix,
                "data_url": f"data:{mime};base64,{b64}",
                "size_bytes": len(data),
            }
        else:
            raise HTTPException(400, f"Unsupported file type: .{suffix}")
    except Exception as e:
        raise HTTPException(500, str(e)[:500])


@app.get("/api/download/{path:path}")
def download_file(path: str):
    """Download a file from the workspace."""
    full_path = Path(_config["workspace"]) / path
    if not full_path.exists():
        raise HTTPException(404, f"File not found: {path}")
    if not str(full_path.resolve()).startswith(str(Path(_config["workspace"]).resolve())):
        raise HTTPException(403, "Access denied")
    return FileResponse(full_path, filename=full_path.name)


@app.post("/api/workspace/move")
def move_file(req: dict):
    """Move or rename a file within the workspace."""
    ws = Path(_config["workspace"])
    src_rel = req.get("src", "").strip()
    dest_dir_rel = req.get("dest_dir", "").strip()
    if not src_rel:
        raise HTTPException(400, "Source path is required")

    src = ws / src_rel
    if not str(src.resolve()).startswith(str(ws.resolve())):
        raise HTTPException(403, "Access denied: source outside workspace")
    if not src.exists():
        raise HTTPException(404, f"File not found: {src_rel}")

    dest_dir = ws / dest_dir_rel if dest_dir_rel else ws
    if not str(dest_dir.resolve()).startswith(str(ws.resolve())):
        raise HTTPException(403, "Access denied: destination outside workspace")
    dest_dir.mkdir(parents=True, exist_ok=True)

    dest = dest_dir / src.name
    if dest.exists():
        raise HTTPException(409, f"Destination already exists: {dest.name}")

    src.rename(dest)
    invalidate_index(str(ws))
    return {"status": "ok", "path": str(dest.relative_to(ws))}


def _rm_readonly(func, path, _):
    """Clear read-only bit and retry deletion (used by shutil.rmtree on Windows)."""
    os.chmod(path, stat.S_IWRITE)
    func(path)


def _try_unlink(path: Path) -> bool:
    """Try to delete a file with a few retries for Windows locking issues."""
    for _ in range(3):
        try:
            path.unlink()
            return True
        except PermissionError:
            time.sleep(0.3)
    return False


@app.post("/api/workspace/delete")
def delete_file(req: dict):
    """Delete a file or folder within the workspace."""
    ws = Path(_config["workspace"])
    rel = req.get("path", "").strip()
    if not rel:
        raise HTTPException(400, "Path is required")

    target = ws / rel
    if not str(target.resolve()).startswith(str(ws.resolve())):
        raise HTTPException(403, "Access denied: path outside workspace")
    if not target.exists():
        raise HTTPException(404, f"Not found: {rel}")

    try:
        if target.is_dir():
            shutil.rmtree(target, onerror=_rm_readonly)
        else:
            if not _try_unlink(target):
                raise PermissionError(
                    f"File is locked by another program. "
                    f"Close apps that may be using '{target.name}' and try again."
                )
        invalidate_index(str(ws))
        return {"status": "ok", "deleted": rel}
    except PermissionError as e:
        raise HTTPException(409, str(e))
    except Exception as e:
        raise HTTPException(500, f"Delete failed: {e}")


# ---------------------------------------------------------------------------
# Serve frontend static files (must be LAST — catch-all)
# ---------------------------------------------------------------------------

_STATIC_DIR = Path(__file__).resolve().parent.parent.parent / "static"


@app.get("/")
async def serve_index():
    return FileResponse(_STATIC_DIR / "index.html")


@app.get("/{full_path:path}")
async def serve_static(full_path: str):
    file_path = _STATIC_DIR / full_path
    if file_path.exists() and file_path.is_file():
        return FileResponse(file_path)
    return FileResponse(_STATIC_DIR / "index.html")
