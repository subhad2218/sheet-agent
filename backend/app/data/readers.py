"""Data readers for common file formats.

Centralises the logic for inspecting Excel, CSV, Parquet and Word
files so it can be reused by both the tool layer and the preview API.
"""

from __future__ import annotations
import json
from pathlib import Path
from typing import Any

from ..sandbox.executor import execute_code


def read_tabular_schema(
    full_path: Path,
    n_rows: int = 5,
    sheet_name: str | None = None,
) -> dict[str, Any]:
    """Read a tabular file and return its schema + sample rows.

    Runs inside the sandbox so the same environment (polars, fastexcel)
    is available as execute_python.
    """
    safe_path = str(full_path).replace("\\", "\\\\")

    sheet_param = ""
    if sheet_name:
        try:
            sheet_id = int(sheet_name)
            sheet_param = f", sheet_id={sheet_id}"
        except ValueError:
            sheet_param = f', sheet_name="{sheet_name}"'

    code = f"""import polars as pl
import json

path = r"{safe_path}"
suffix = path.lower().split('.')[-1]

if suffix == 'csv':
    df = pl.read_csv(path, infer_schema_length=200, ignore_errors=True)
elif suffix == 'parquet':
    df = pl.read_parquet(path)
else:
    df = pl.read_excel(path{sheet_param})

schema = {{col: str(dtype) for col, dtype in zip(df.columns, df.dtypes)}}
sample = df.head({n_rows}).to_dicts()

result = {{"columns": list(df.columns), "schema": schema, "row_count": df.height, "sample": sample}}
print(json.dumps(result, ensure_ascii=False, default=str))
"""
    workspace = str(full_path.parent.parent) if full_path.parent.name == ".cache" else str(full_path.parent)
    r = execute_code(code, workspace, cwd=str(full_path.parent))
    if r.success:
        return json.loads(r.stdout.strip())
    raise RuntimeError(r.stderr[:500])


def read_document(full_path: Path, max_chars: int = 5000) -> dict[str, Any]:
    """Read a Word document and return text + tables."""
    safe_path = str(full_path).replace("\\", "\\\\")

    code = f"""import json
import os
path = r"{safe_path}"
suffix = path.lower().split(".")[-1]

text_parts = []
tables = []

if suffix == "docx":
    try:
        from docx import Document
        doc = Document(path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            tbl = []
            for row in table.rows:
                tbl.append([cell.text for cell in row.cells])
            tables.append(tbl)
    except Exception as e:
        text_parts.append(f"Error reading docx: {{str(e)}}")
elif suffix == "doc":
    try:
        import win32com.client as wc
        word = wc.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(path))
        text_parts.append(doc.Content.Text)
        doc.Close(False)
        word.Quit()
    except Exception as e:
        text_parts.append(f"Error reading doc (try converting to docx): {{str(e)}}")
else:
    text_parts.append(f"Unsupported document format: {{suffix}}")

full_text = "\\n".join(text_parts)
result = {{
    "path": path,
    "format": suffix,
    "text_length": len(full_text),
    "text_preview": full_text[:{max_chars}],
    "paragraph_count": len(text_parts),
    "tables_found": len(tables),
    "tables": tables[:5],
}}
print(json.dumps(result, ensure_ascii=False))
"""
    workspace = str(full_path.parent.parent) if full_path.parent.name == ".cache" else str(full_path.parent)
    r = execute_code(code, workspace, cwd=str(full_path.parent))
    if r.success:
        return json.loads(r.stdout.strip())
    raise RuntimeError(r.stderr[:500])
