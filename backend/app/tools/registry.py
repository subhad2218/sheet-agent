"""Tool definitions and execution dispatcher."""

from __future__ import annotations
import json
import os
from pathlib import Path
from typing import Any

from ..llm.base import ToolDef, ToolCall
from ..sandbox.executor import execute_code, ExecutionResult


# ---------------------------------------------------------------------------
# Tool definitions — provider-agnostic
# ---------------------------------------------------------------------------

TOOLS: list[ToolDef] = [
    ToolDef(
        name="list_files",
        description="List files in the workspace. Use glob patterns like '*.xlsx' to filter.",
        parameters={
            "pattern": {
                "type": "string",
                "description": "Glob pattern to match files, e.g. '*.xlsx', '**/*.csv'",
                "required": False,
            },
        },
    ),
    ToolDef(
        name="read_excel",
        description="Read an Excel, CSV, or Parquet file and return its schema (column names, dtypes) and first N rows as sample data. For multi-sheet Excel files, use sheet_name to read a specific sheet.",
        parameters={
            "path": {
                "type": "string",
                "description": "Relative path to the file within the workspace",
                "required": True,
            },
            "sheet_name": {
                "type": "string",
                "description": "Sheet name (or 1-based index) to read for multi-sheet Excel files. Leave empty to read the first sheet.",
                "required": False,
            },
            "n_rows": {
                "type": "integer",
                "description": "Number of sample rows to return (default 5)",
                "required": False,
            },
        },
    ),
    ToolDef(
        name="read_document",
        description="Read a Word document (.docx or .doc) and return its text content and tables. Use this to inspect documents before processing them.",
        parameters={
            "path": {
                "type": "string",
                "description": "Relative path to the document within the workspace",
                "required": True,
            },
            "max_chars": {
                "type": "integer",
                "description": "Maximum characters of text to return (default 5000)",
                "required": False,
            },
        },
    ),
    ToolDef(
        name="execute_python",
        description="Execute Python code to process data. The 'polars' library is available as 'pl'. Use pl.read_excel()/pl.read_csv()/pl.scan_parquet() to read files, and pl.DataFrame.write_excel()/write_csv() to write results. The current working directory is the workspace.",
        parameters={
            "code": {
                "type": "string",
                "description": "Python code to execute",
                "required": True,
            },
        },
    ),
    ToolDef(
        name="write_result",
        description="Export a file that was created by execute_python. Use this after your code has produced an output file to confirm and register the result.",
        parameters={
            "path": {
                "type": "string",
                "description": "Relative path of the output file within workspace",
                "required": True,
            },
            "format": {
                "type": "string",
                "description": "Output format",
                "enum": ["xlsx", "csv", "parquet"],
                "required": False,
            },
        },
    ),
]


# ---------------------------------------------------------------------------
# Tool execution
# ---------------------------------------------------------------------------

def _list_files(workspace: str, current_dir: str, pattern: str = "*") -> str:
    target = Path(workspace) / current_dir if current_dir else Path(workspace)
    matches = sorted(target.glob(pattern))
    if not matches:
        return json.dumps({"files": [], "message": f"No files matching '{pattern}' in {current_dir or '/'}"})

    result = []
    for f in matches:
        if f.is_file():
            stat = f.stat()
            result.append({
                "name": f.name,
                "path": str(f.relative_to(Path(workspace))),
                "size_bytes": stat.st_size,
                "size_mb": round(stat.st_size / 1024 / 1024, 2),
            })
    return json.dumps({"files": result, "count": len(result)}, ensure_ascii=False)


def _read_excel(workspace: str, current_dir: str, path: str, n_rows: int = 5, sheet_name: str | None = None) -> str:
    """Read file via sandbox — same environment as execute_python."""
    target_dir = Path(workspace) / current_dir if current_dir else Path(workspace)
    full_path = target_dir / path
    if not full_path.exists():
        return json.dumps({"error": f"File not found: {path} (in {current_dir or '/'}))"}, ensure_ascii=False)

    # Use raw string for Windows paths, escape backslashes for safety
    safe_path = str(full_path).replace("\\", "\\\\")

    # Build sheet parameter for polars.read_excel
    sheet_param = ""
    if sheet_name:
        try:
            sheet_id = int(sheet_name)
            sheet_param = f", sheet_id={sheet_id}"
        except ValueError:
            sheet_param = f', sheet_name="{sheet_name}"'

    code = f"""import polars as pl
import json

path = r"{safe_path}"
suffix = path.lower().split('.')[-1]

if suffix == 'csv':
    df = pl.read_csv(path, infer_schema_length=200, ignore_errors=True)
elif suffix == 'parquet':
    df = pl.read_parquet(path)
else:
    df = pl.read_excel(path{sheet_param})

schema = {{col: str(dtype) for col, dtype in zip(df.columns, df.dtypes)}}
sample = df.head({n_rows}).to_dicts()

result = {{"columns": list(df.columns), "schema": schema, "row_count": df.height, "sample": sample}}
print(json.dumps(result, ensure_ascii=False, default=str))
"""
    r = execute_code(code, workspace, cwd=str(target_dir))
    if r.success:
        return r.stdout.strip()
    else:
        return json.dumps({"error": r.stderr[:500], "stdout": r.stdout[:200]}, ensure_ascii=False)


def _read_document(workspace: str, current_dir: str, path: str, max_chars: int = 5000) -> str:
    """Read Word document via sandbox and return text + tables."""
    target_dir = Path(workspace) / current_dir if current_dir else Path(workspace)
    full_path = target_dir / path
    if not full_path.exists():
        return json.dumps({"error": f"File not found: {path} (in {current_dir or '/'})"}, ensure_ascii=False)

    safe_path = str(full_path).replace("\\", "\\\\")

    code = f"""import json
import os
path = r"{safe_path}"
suffix = path.lower().split(".")[-1]

text_parts = []
tables = []

if suffix == "docx":
    try:
        from docx import Document
        doc = Document(path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            tbl = []
            for row in table.rows:
                tbl.append([cell.text for cell in row.cells])
            tables.append(tbl)
    except Exception as e:
        text_parts.append(f"Error reading docx: {{str(e)}}")
elif suffix == "doc":
    try:
        import win32com.client as wc
        word = wc.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(path))
        text_parts.append(doc.Content.Text)
        doc.Close(False)
        word.Quit()
    except Exception as e:
        text_parts.append(f"Error reading doc (try converting to docx): {{str(e)}}")
else:
    text_parts.append(f"Unsupported document format: {{suffix}}")

full_text = "\\n".join(text_parts)
result = {{
    "path": path,
    "format": suffix,
    "text_length": len(full_text),
    "text_preview": full_text[:{max_chars}],
    "paragraph_count": len(text_parts),
    "tables_found": len(tables),
    "tables": tables[:5],
}}
print(json.dumps(result, ensure_ascii=False))
"""
    r = execute_code(code, workspace, cwd=str(target_dir))
    if r.success:
        return r.stdout.strip()
    else:
        return json.dumps({"error": r.stderr[:500], "stdout": r.stdout[:200]}, ensure_ascii=False)


def _execute_python(workspace: str, current_dir: str, code: str) -> str:
    target_dir = Path(workspace) / current_dir if current_dir else Path(workspace)
    r = execute_code(code, workspace, cwd=str(target_dir))
    result = {
        "success": r.success,
        "stdout": r.stdout[-2000:] if r.stdout else "",
        "stderr": r.stderr[-2000:] if r.stderr else "",
        "duration_ms": r.duration_ms,
    }
    if not r.success:
        stderr = r.stderr.lower()
        hint = "Code execution failed. Review the error and try again."
        if "column" in stderr or "keyerror" in stderr or "not found" in stderr:
            hint += " Common cause: column name mismatch. Use read_excel to verify exact column names (case-sensitive)."
        elif "encoding" in stderr or "unicode" in stderr:
            hint += " Common cause: encoding issue. Try encoding='utf-8-sig' or encoding='gbk' for CSV."
        elif "permission" in stderr or "access denied" in stderr:
            hint += " Common cause: file permission or path outside workspace. Use relative paths only."
        elif "memory" in stderr or "out of" in stderr:
            hint += " Common cause: out of memory. Process data in smaller chunks or filter rows first."
        elif "module" in stderr or "importerror" in stderr or "no module" in stderr:
            hint += " Common cause: unsupported import. Only polars, matplotlib, seaborn, docx, and standard library are available."
        elif "sheet" in stderr:
            hint += " Common cause: invalid sheet name or index. Use read_excel to list available sheets."
        else:
            hint += " Common fixes: check column names with read_excel first, handle encoding issues, verify file paths."
        result["hint"] = hint
    return json.dumps(result, ensure_ascii=False)


def _write_result(workspace: str, current_dir: str, path: str, format: str = "xlsx") -> str:
    target_dir = Path(workspace) / current_dir if current_dir else Path(workspace)
    full_path = target_dir / path
    if not full_path.exists():
        return json.dumps({"error": f"File not found: {path} (in {current_dir or '/'})"}, ensure_ascii=False)

    stat = full_path.stat()
    return json.dumps({
        "success": True,
        "path": str(full_path.relative_to(Path(workspace))),
        "format": format,
        "size_bytes": stat.st_size,
        "size_mb": round(stat.st_size / 1024 / 1024, 2),
    }, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def execute_tool(call: ToolCall, workspace: str, current_dir: str = "") -> str:
    """Execute a tool call and return the result as a JSON string."""
    name = call.name
    args = call.arguments
    try:
        if name == "list_files":
            return _list_files(workspace, current_dir, args.get("pattern", "*"))
        elif name == "read_excel":
            return _read_excel(workspace, current_dir, args["path"], args.get("n_rows", 5), args.get("sheet_name"))
        elif name == "read_document":
            return _read_document(workspace, current_dir, args["path"], args.get("max_chars", 5000))
        elif name == "execute_python":
            return _execute_python(workspace, current_dir, args["code"])
        elif name == "write_result":
            return _write_result(workspace, current_dir, args["path"], args.get("format", "xlsx"))
        else:
            return json.dumps({"error": f"Unknown tool: {name}"})
    except KeyError as e:
        return json.dumps({"success": False, "error": f"Missing required parameter: {e}. The tool '{name}' requires this parameter. Please check your tool call arguments.", "stderr": f"Missing required parameter: {e}"})
